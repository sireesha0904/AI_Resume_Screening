from docx import Document
import os
def extract_text_from_resume(file_path):
    """Extracts text from a .docx resume file given its path."""
    try:
        document = Document(file_path)

        resume_text = []
        # Extract paragraphs
        for para in document.paragraphs:
            resume_text.append(para.text)

        # Extract table data (if any)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    resume_text.append(cell.text)

        extracted_text = ' '.join(resume_text).strip()
        print(f"Extracted text from {file_path}:", extracted_text[:300])  # Optional debug
        return extracted_text
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return ""